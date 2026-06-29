"""
agents/writer_agent.py  —  Writer Agent (업그레이드)

기능:
  1. 에세이/문서 첨삭 및 개선 (기존)
  2. 웹 서핑 후 주제 조사 → Word (.docx) 파일로 저장 (신규)
  3. 이메일·커버레터·보고서 초안 작성 (기존)

docx 저장 위치: Agent J/outputs/ 폴더

사용 예시:
  "AI 에이전트 동향 조사 후 워드 파일로 정리해줘"
  "이 에세이 첨삭해줘: [내용]"
  "데이터마이닝 기초 개념 정리 워드로 만들어줘"
"""

import os
import re
from datetime import date, datetime
from pathlib import Path
from agents.base_agent import BaseAgent
from tools.file_tools import DEV_TOOLS, execute_tool

OUTPUTS_DIR = Path(__file__).parent.parent / "outputs"

DOCX_KEYWORDS = [
    "워드", "word", ".docx", "docx", "파일로", "문서로 저장", "정리해줘",
    "조사 후", "웹서핑", "리서치 후", "보고서로", "문서 만들어", "저장해줘"
]

RESEARCH_KEYWORDS = [
    "조사", "리서치", "검색", "알아봐", "찾아봐", "웹서핑", "최신", "동향", "트렌드", "정리해줘"
]

WRITER_SYSTEM_PROMPT = """You are J's Writer Agent — a professional editor, writing coach, and research-to-document specialist.

## Roles
1. **Essay & Document Editing**: structure, flow, argument clarity, transitions
2. **Web Research → Document**: synthesize research findings into a structured Word document
3. **Drafting**: emails, reports, cover letters, summaries
4. **Translation**: Korean ↔ English

## Principles
- Explain WHAT you changed and WHY (briefly).
- Offer 2-3 alternative phrasings for key sentences when relevant.
- Preserve the author's voice — don't over-sanitize.
- For research documents: use ## headers for sections, cite sources at the end.
- Final output must be publish-ready.

## Research Document Format (when asked to research + save)
# [주제]
## 개요
## 핵심 내용
### [섹션 1]
### [섹션 2]
## 주요 개념
## 참고 자료

Today's date: {today}
"""


def _has_docx_intent(message: str) -> bool:
    msg = message.lower()
    return any(kw in msg for kw in DOCX_KEYWORDS)


def _has_research_intent(message: str) -> bool:
    msg = message.lower()
    return any(kw in msg for kw in RESEARCH_KEYWORDS)


def _web_research(topic: str) -> str:
    """주제를 웹 검색해서 컨텍스트 텍스트로 반환."""
    try:
        from tools.research_tools import web_search, fetch_page_text
        print(f"  🔍 '{topic[:50]}' 웹 조사 중...")
        results = web_search(topic, num_results=5)
        if not results:
            return ""
        lines = []
        for i, r in enumerate(results, 1):
            lines.append(
                f"[출처 {i}] {r.get('title', '')}\n"
                f"{r.get('snippet', '')}\nURL: {r.get('url', '')}"
            )
            if i == 1 and r.get("url"):
                page = fetch_page_text(r["url"], max_chars=1500)
                if page:
                    lines.append(f"(상세 내용): {page[:800]}")
        return "\n\n".join(lines)
    except Exception as e:
        return f"(웹 조사 오류: {e})"


def _save_docx(title: str, content: str) -> dict:
    """
    python-docx로 .docx 파일을 outputs/ 폴더에 저장.
    python-docx 없으면 .txt 폴백.
    """
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r'[\\/*?:"<>|]', "_", title)[:50]
    timestamp  = datetime.now().strftime("%Y%m%d_%H%M")
    filename   = f"{safe_title}_{timestamp}"

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc     = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(f"생성일: {date.today().isoformat()} | Agent J")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith("#### "):
                doc.add_heading(stripped[5:], level=3)
            elif stripped.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s*", "", stripped), style="List Number")
            elif stripped:
                doc.add_paragraph(stripped)

        path = OUTPUTS_DIR / f"{filename}.docx"
        doc.save(str(path))
        return {"success": True, "path": str(path), "format": "docx"}

    except ImportError:
        path = OUTPUTS_DIR / f"{filename}.txt"
        path.write_text(f"{title}\n{'='*len(title)}\n\n{content}", encoding="utf-8")
        return {
            "success": True, "path": str(path), "format": "txt",
            "note": "`pip install python-docx` 설치 후 재실행하면 .docx로 저장돼요."
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


class WriterAgent(BaseAgent):

    def __init__(self):
        model  = os.getenv("DEV_MODEL", "claude-sonnet-4-6")
        system = WRITER_SYSTEM_PROMPT.format(today=date.today().isoformat())
        super().__init__(
            model=model,
            system_prompt=system,
            tools=DEV_TOOLS,
            tool_executor=execute_tool,
            name="Writer Agent"
        )

    def run(self, user_message: str, history: list = None) -> str:
        save_docx   = _has_docx_intent(user_message)
        do_research = _has_research_intent(user_message)

        augmented_message = user_message
        if do_research:
            web_ctx = _web_research(user_message[:120])
            if web_ctx:
                augmented_message = (
                    f"{user_message}\n\n"
                    f"[웹 조사 결과 — 아래 내용을 바탕으로 문서를 작성해주세요]\n{web_ctx}"
                )

        response = super().run(augmented_message, history)

        if save_docx:
            title_match = re.search(r'^#+ (.+)', response, re.MULTILINE)
            title = title_match.group(1) if title_match else user_message[:40].strip()
            print(f"  💾 저장 중: outputs/{title[:30]}...")
            save_result = _save_docx(title, response)

            if save_result["success"]:
                path = save_result["path"]
                fmt  = save_result["format"]
                note = save_result.get("note", "")
                suffix = f"\n\n---\n✅ **{fmt.upper()} 저장 완료**\n📁 `{path}`"
                if note:
                    suffix += f"\n⚠️ {note}"
                response += suffix
            else:
                response += f"\n\n⚠️ 파일 저장 실패: {save_result.get('error')}"

        return response
